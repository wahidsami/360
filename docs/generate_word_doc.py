"""
Arena360 Word Document Generator
Generates a professional .docx presentation guide
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Color palette ────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x23, 0x3B)   # headings
ACCENT     = RGBColor(0x4F, 0x46, 0xE5)   # accent / links
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)   # table header bg
MID_GRAY   = RGBColor(0x6B, 0x72, 0x80)   # body text
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GREEN      = RGBColor(0x16, 0xA3, 0x4A)
RED        = RGBColor(0xDC, 0x26, 0x26)

doc = Document()

# ─── Page margins ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ─── Helper: set paragraph background via w:shd ───────────────────
def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),  kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'),   kwargs[edge].get('sz', '4'))
            tag.set(qn('w:color'),kwargs[edge].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

# ─── Helper: add a heading ────────────────────────────────────────
def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = DARK_BLUE
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(15)
    else:
        run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

# ─── Helper: add body paragraph ───────────────────────────────────
def add_para(text, bold=False, italic=False, size=10.5, color=None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else MID_GRAY
    p.paragraph_format.space_after = Pt(4)
    return p

# ─── Helper: add bullet ───────────────────────────────────────────
def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = MID_GRAY
    return p

# ─── Helper: add table with header ───────────────────────────────
def add_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    table  = doc.add_table(rows=1, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        shade_cell(cell, '1A233B')
        p   = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE
        p.alignment        = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cell = row_cells[i]
            p    = cell.paragraphs[0]
            run  = p.add_run(str(val))
            run.font.size      = Pt(9)
            run.font.color.rgb = MID_GRAY
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if str(val) in ('✅','❌','✔','✗') else WD_ALIGN_PARAGRAPH.LEFT
    # col widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()  # spacer
    return table

# ═══════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ═══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ARENA360')
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Project Management Platform')
r.font.size = Pt(18)
r.font.color.rgb = ACCENT

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Comprehensive Presentation & Demo Guide')
r.font.size = Pt(14)
r.font.color.rgb = MID_GRAY
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Version 1.0  |  March 2026')
r.font.size = Pt(11)
r.font.color.rgb = MID_GRAY

doc.add_paragraph('\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('For Project Managers, Operations Teams & Client Presentations')
r.font.size = Pt(12)
r.bold = True
r.font.color.rgb = ACCENT

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  §1  WHAT IS ARENA360?
# ═══════════════════════════════════════════════════════════════════
add_heading('1. What is Arena360?', level=1)
add_para(
    'Arena360 is a full-stack, multi-tenant project management and operations platform '
    'built for digital agencies, IT service companies, and consultancy firms. It acts as '
    'a centralized command center where internal teams and clients collaborate in one '
    'unified, role-aware interface.',
    size=11
)

add_heading('Core Value Proposition', level=3)
add_para(
    '"One platform for your entire project delivery lifecycle — '
    'from client onboarding to invoice collection."',
    italic=True, bold=True, size=11, color=ACCENT
)

add_heading('What Arena360 Includes', level=3)
features = [
    'Project Management — tasks, sprints, milestones, Gantt, Kanban',
    'Client Portal — external-facing, role-scoped client access',
    'Financial Management — contracts, invoices, Stripe payments',
    'QA / Findings Tracking — severity levels, evidence, AI analysis',
    'Real-Time Notifications — Socket.IO WebSocket push',
    'Workflow Automation — rule-based triggers, multi-step approvals',
    'Report Generation — PPTX and PDF from live project data',
    'Time Tracking — billable/non-billable hours per task',
    'Arabic / RTL Support — MENA-region ready',
    'AI Features (optional) — project summaries, task suggestions, chat',
]
for f in features: add_bullet(f)
doc.add_paragraph()

add_heading('Target Users', level=3)
add_table(
    ['Who', 'What They Do on Arena360'],
    [
        ['Project Managers (PM)', 'Create projects, assign tasks, track milestones, manage team, generate reports'],
        ['Developers (DEV)',       'See assigned tasks, log time, update statuses, respond to findings'],
        ['Finance Teams',          'Manage contracts, invoices, financial KPIs, approve payments'],
        ['Operations (OPS)',        'Oversee all clients, projects, financials, and org health'],
        ['Clients',                'Monitor their projects, view updates, download files, pay invoices'],
        ['Super Admins',           'Control the entire platform — users, SSO, branding, audit logs'],
    ],
    col_widths=[2.2, 4.5]
)

# ═══════════════════════════════════════════════════════════════════
#  §2  TECHNOLOGY ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════
add_heading('2. Technology Architecture', level=1)

add_heading('Frontend Stack', level=2)
add_table(
    ['Technology', 'Purpose'],
    [
        ['React 18 + TypeScript',  'Single-page application'],
        ['Vite',                   'Fast build & dev server'],
        ['React Router v6',        'Client-side navigation'],
        ['Tailwind CSS',           'Styling & responsive design'],
        ['Recharts',               'Charts and analytics dashboards'],
        ['FullCalendar',           'Task/milestone calendar view'],
        ['react-frappe-gantt',     'Gantt timeline view'],
        ['react-i18next',          'English / Arabic, RTL support'],
        ['Socket.IO Client',       'Real-time WebSocket notifications'],
    ],
    col_widths=[2.5, 4.2]
)

add_heading('Backend Stack', level=2)
add_table(
    ['Technology', 'Purpose'],
    [
        ['NestJS (Node.js)',       'REST API server'],
        ['Prisma ORM',            'Type-safe database access layer'],
        ['PostgreSQL 15+',        'Relational data store'],
        ['JWT + bcrypt + TOTP',   'Auth, sessions, 2FA'],
        ['Google OAuth + SAML',   'Enterprise SSO'],
        ['Socket.IO',             'Real-time WebSocket layer'],
        ['Stripe',                'Invoice payment processing'],
        ['Resend',                'Transactional email (invites, resets)'],
        ['pptxgenjs + pdfkit',    'Report generation (PPTX / PDF)'],
        ['MinIO (S3-compatible)', 'File and document storage'],
        ['OpenAI API (optional)', 'AI summaries, suggestions, chat'],
        ['Swagger / OpenAPI',     'Interactive API docs at /api-docs'],
    ],
    col_widths=[2.5, 4.2]
)

# ═══════════════════════════════════════════════════════════════════
#  §3  USER TYPES & ROLES
# ═══════════════════════════════════════════════════════════════════
add_heading('3. User Types & Roles', level=1)
add_para(
    'Arena360 uses a 9-role permission model divided into Internal (Staff) '
    'and External (Client) personas. Every role sees a tailored interface.',
    size=11
)

add_heading('Internal Roles (Staff)', level=2)
add_table(
    ['Role', 'Title', 'Key Permissions'],
    [
        ['SUPER_ADMIN', 'System Administrator', 'All permissions — users, SSO, branding, audit, admin panel'],
        ['OPS',         'Operations Manager',   'All clients, projects, financials, tasks, team management'],
        ['PM',          'Project Manager',       'Projects, tasks, team — no financial management'],
        ['DEV',         'Developer',             'Dashboard, view clients, manage their own assigned tasks'],
        ['FINANCE',     'Financial Officer',     'Dashboard, financials, clients — read-heavy role'],
    ],
    col_widths=[1.5, 1.8, 3.4]
)

add_heading('External Roles (Client Portal)', level=2)
add_table(
    ['Role', 'Title', 'What They Can See'],
    [
        ['CLIENT_OWNER',   'Client Organization Owner', 'Projects, client-visible updates, financials, pay invoices'],
        ['CLIENT_MANAGER', 'Client Team Manager',       'Projects and client-scoped data'],
        ['CLIENT_MEMBER',  'Regular Client Member',     'Dashboard only'],
        ['VIEWER',         'Read-Only Guest',           'Dashboard only'],
    ],
    col_widths=[1.5, 1.8, 3.4]
)

add_heading('Permission Matrix', level=2)
add_table(
    ['Permission', 'SUPER_ADMIN', 'OPS', 'PM', 'DEV', 'FINANCE', 'CLIENT_OWNER', 'CLIENT_MANAGER', 'CLIENT_MEMBER', 'VIEWER'],
    [
        ['VIEW_DASHBOARD',  '✅','✅','✅','✅','✅','✅','✅','✅','✅'],
        ['VIEW_CLIENTS',    '✅','✅','✅','✅','✅','✅','✅','❌','❌'],
        ['MANAGE_CLIENTS',  '✅','✅','❌','❌','❌','❌','❌','❌','❌'],
        ['MANAGE_PROJECTS', '✅','✅','✅','❌','❌','❌','❌','❌','❌'],
        ['MANAGE_TASKS',    '✅','✅','✅','✅','❌','❌','❌','❌','❌'],
        ['MANAGE_TEAM',     '✅','✅','✅','❌','❌','❌','❌','❌','❌'],
        ['VIEW_FINANCIALS', '✅','✅','❌','❌','✅','✅','❌','❌','❌'],
        ['MANAGE_USERS',    '✅','❌','❌','❌','❌','❌','❌','❌','❌'],
        ['VIEW_ADMIN',      '✅','❌','❌','❌','❌','❌','❌','❌','❌'],
    ],
    col_widths=[1.7, 0.7, 0.5, 0.5, 0.5, 0.7, 0.9, 1.0, 0.9, 0.65]
)

# ═══════════════════════════════════════════════════════════════════
#  §4  ROLE-BASED DASHBOARDS
# ═══════════════════════════════════════════════════════════════════
add_heading('4. Role-Based Dashboards', level=1)
add_para(
    'One of Arena360\'s strongest differentiators — every persona sees a tailored command center. '
    'No information overload, no confusion.',
    size=11
)

dashboards = [
    (
        '4.1 Admin Dashboard (SUPER_ADMIN / OPS / PM)',
        'The operational overview for managers. Ideal for: daily standups, executive reporting, portfolio health checks.',
        [
            'KPI Cards: Total Clients | Active Projects | Revenue (SAR) | Overdue Tasks (with trend arrows)',
            'Revenue Velocity Chart: Area chart showing monthly revenue trends',
            'Latest Project Updates: Cross-project feed of newest status posts',
            'Projects at Risk: List of projects flagged as AT_RISK or CRITICAL',
            'Pending Approvals: Count of items awaiting manager review',
            'Quick Actions / Tools Panel: Shortcuts to create clients, projects, tasks',
            'Dashboard Customization: Show/hide and reorder widgets (stored per-user)',
        ]
    ),
    (
        '4.2 Developer Dashboard (DEV)',
        'The personal task command center. Ideal for: morning task review, end-of-day progress check.',
        [
            'My Open Tasks: All tasks assigned to this user with status badges',
            'Due Soon: Tasks with deadlines within 3 days (amber warning)',
            'In Review: Tasks currently awaiting review from manager',
            'Overdue Counter: Tasks past their deadline (red alert)',
            'Project Quick Navigation: One-click jump to parent project',
        ]
    ),
    (
        '4.3 Finance Dashboard (FINANCE)',
        'The financial health overview. Ideal for: cash flow management, financial reporting.',
        [
            'Total Revenue: Paid invoices across all projects',
            'Outstanding Balance: Unpaid / overdue invoices',
            'Invoices Due: Invoices approaching due date',
            'Revenue Breakdown Charts: Monthly financial trend visualization',
        ]
    ),
    (
        '4.4 Client Dashboard (CLIENT_*)',
        'The client portal view — clients see only their own data. Ideal for: client check-ins, status reporting.',
        [
            'My Projects: All projects linked to the client\'s organization',
            'Project Status: Aggregated health and progress per project',
            'Recent Updates: Status posts marked as CLIENT-visible',
            'Shared Files: Files uploaded with Client visibility only',
            'Financials (CLIENT_OWNER only): Contracts and invoices; Pay with Card button',
        ]
    ),
]

for title, desc, bullets in dashboards:
    add_heading(title, level=2)
    add_para(desc, size=10.5)
    for b in bullets: add_bullet(b)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
#  §5  CORE FEATURES
# ═══════════════════════════════════════════════════════════════════
add_heading('5. Core Features Overview', level=1)

add_heading('5.1 Project Management — 15 Specialized Tabs', level=2)
add_table(
    ['Tab', 'What It Contains'],
    [
        ['Overview',        'Description, status, health indicator, progress %, dates, budget, tags'],
        ['Tasks (Kanban)',  'Backlog → Todo → In Progress → Review → Done; assign, priority, due date'],
        ['Time',            'Time entries per task (date, minutes, billable/non-billable, note)'],
        ['Timeline (Gantt)','Visual project timeline with tasks and dependencies'],
        ['Sprints',         'Sprint creation, backlog management, assign tasks to sprints'],
        ['Recurring Tasks', 'Task templates with cron-based scheduling; enable/disable'],
        ['Milestones',      'Milestone tracking with % complete, owner, due date'],
        ['Updates',         'Status updates with Internal or Client visibility'],
        ['Files',           'Upload/download files by category and visibility scope'],
        ['Findings',        'Link to QA/issue tracking module'],
        ['Reports',         'Generate PPTX/PDF reports from live project data'],
        ['Financials',      'Contracts, invoices, Stripe "Pay with Card" payment'],
        ['Team',            'Add/remove members, assign project roles'],
        ['Discussions',     'Threaded project-level discussions and replies'],
        ['Activity',        'Full activity feed / audit trail for the project'],
    ],
    col_widths=[1.8, 4.9]
)

add_heading('5.2 Task Management', level=2)
add_table(
    ['Feature', 'Detail'],
    [
        ['Workflow',       'BACKLOG → TODO → IN_PROGRESS → REVIEW → DONE'],
        ['Fields',         'Title, description, priority, assignee, due date, labels, milestone, sprint, story points'],
        ['Dependencies',   'Predecessor/successor task links'],
        ['Recurring',      'Auto-created from templates via cron scheduler'],
        ['My Work View',   'Personal task summary with KPI cards (Open, Due Soon, In Review, Overdue)'],
        ['Priority Colors','Urgent (red), High (amber), Normal (cyan)'],
    ],
    col_widths=[2.0, 4.7]
)

add_heading('5.3 QA / Findings Module', level=2)
add_table(
    ['Aspect', 'Detail'],
    [
        ['Severity Levels', 'Low / Medium / High / Critical'],
        ['Status Workflow',  'Open → In Progress → In Review → Closed / Dismissed'],
        ['Evidence',         'File uploads — screenshots, logs, test results'],
        ['Comments',         'Threaded comment system with nested replies'],
        ['Timeline',         'Full activity history per finding'],
        ['AI Analysis',      'Optional — OpenAI suggests remediation and confirms severity'],
        ['GitHub Integration','Create GitHub issues directly from findings'],
    ],
    col_widths=[2.0, 4.7]
)

add_heading('5.4 Financial Management', level=2)
add_table(
    ['Feature', 'Detail'],
    [
        ['Contracts',         'Amount, currency (SAR default), dates, Active/Completed/Cancelled'],
        ['Invoices',          'Invoice number, amount, due date, Draft/Issued/Paid/Overdue'],
        ['Stripe Payments',   '"Pay with Card" via Stripe Elements; webhook marks invoice Paid automatically'],
        ['Multi-step Approvals','Approval chains for reports, invoices, and contracts'],
        ['Access Control',    'Only SUPER_ADMIN, OPS, FINANCE, CLIENT_OWNER can access financials'],
    ],
    col_widths=[2.0, 4.7]
)

add_heading('5.5 Reporting & Analytics', level=2)
add_table(
    ['Feature', 'Detail'],
    [
        ['Report Types',        'Technical, Executive, Compliance, Security, Financial, Status'],
        ['Export Formats',      'PPTX (PowerPoint) and PDF — from live project data'],
        ['CSV Exports',         'Tasks, findings, invoices via export buttons on list views'],
        ['Analytics Page',      'Velocity (tasks/week), completion rate, portfolio charts'],
        ['Project Analytics',   'Team performance, task breakdown, health trends'],
    ],
    col_widths=[2.0, 4.7]
)

# ═══════════════════════════════════════════════════════════════════
#  §6  HOW PM & TEAM USES ARENA360
# ═══════════════════════════════════════════════════════════════════
add_heading('6. How a PM & Team Uses Arena360', level=1)

add_heading('Complete Project Lifecycle End-to-End', level=2)
phases = [
    ('Phase 1 — SETUP', [
        'Admin creates Client account (CRM module)',
        'PM creates Project linked to that Client',
        'PM sets project budget, dates, and tags',
        'PM adds team members (developers, QA)',
    ]),
    ('Phase 2 — PLANNING', [
        'PM creates Milestones (major deliverable checkpoints)',
        'PM creates Sprints if using Agile methodology',
        'PM creates Tasks, assigns to Milestones and Sprints',
        'PM sets up recurring task templates for routine work',
    ]),
    ('Phase 3 — EXECUTION', [
        'Developers see their tasks on dashboard & My Work page',
        'Team moves tasks: TODO → IN_PROGRESS → REVIEW → DONE',
        'Team logs time entries (billable/non-billable) per task',
        'PM posts Internal status updates for team awareness',
        'Automation rules fire notifications on status changes',
    ]),
    ('Phase 4 — QA & REVIEW', [
        'QA/PM creates Findings (bugs, risks, issues)',
        'Findings assigned to developers with severity levels',
        'Evidence uploaded (screenshots, logs, test results)',
        'Findings closed when fixed and reviewed',
    ]),
    ('Phase 5 — CLIENT REPORTING', [
        'PM posts Client-visible Updates via Updates tab',
        'PM shares files with Client visibility',
        'PM generates PPTX/PDF Report via Reports tab',
        'Report goes through multi-step Approval workflow',
        'Client receives notification and sees update in portal',
    ]),
    ('Phase 6 — FINANCIAL CLOSE', [
        'Contract created in Financials tab',
        'Invoices issued against contract',
        'Client sees invoices in their portal',
        'Client pays via Stripe "Pay with Card"',
        'Webhook marks Invoice as PAID automatically',
        'Project status updated to COMPLETED',
    ]),
]

for phase_title, steps in phases:
    add_heading(phase_title, level=3)
    for s in steps: add_bullet(s)
    doc.add_paragraph()

add_heading('Daily Workflows by Role', level=2)
add_table(
    ['Role', 'Daily Workflow in Arena360'],
    [
        ['PM / Admin',
         '1. Check Admin Dashboard (KPIs, at-risk projects)\n'
         '2. Review Pending Approvals\n'
         '3. Open Project → check Kanban board\n'
         '4. Assign or reassign tasks\n'
         '5. Post a Project Update\n'
         '6. Review Findings from QA\n'
         '7. Check Financial tab for invoice statuses'],
        ['Developer',
         '1. Open Developer Dashboard → My Tasks list\n'
         '2. Check Due Soon alerts (amber warning)\n'
         '3. Open a task → move to IN_PROGRESS\n'
         '4. Log time worked on the task\n'
         '5. Submit task for REVIEW when done\n'
         '6. Respond to Finding comments if assigned'],
        ['Finance',
         '1. Open Finance Dashboard → KPI cards\n'
         '2. Review outstanding invoices\n'
         '3. Check contracts per project\n'
         '4. Approve financial documents in workflow'],
        ['Client (CLIENT_OWNER)',
         '1. Log in → Client Dashboard\n'
         '2. View project status and health\n'
         '3. Read client-visible updates\n'
         '4. Download shared files\n'
         '5. View and pay invoices online'],
    ],
    col_widths=[1.5, 5.2]
)

# ═══════════════════════════════════════════════════════════════════
#  §7  DEMO FLOW
# ═══════════════════════════════════════════════════════════════════
add_heading('7. Step-by-Step Demo Flow', level=1)
add_para(
    'Use this guide to walk through the live demonstration for clients and stakeholders. '
    'Demo scenario: "Acme Corp Website Redesign Project"',
    size=11, bold=True
)

demo_steps = [
    ('Step 1 — Login & Admin Dashboard (2 min)',
     'Login as PM. Show the Admin Dashboard: KPI cards (Clients, Projects, Revenue, Overdue Tasks), '
     'Revenue Velocity chart, Projects at Risk list, Pending Approvals. Say: '
     '"As a PM, this is my command center — everything I need at a glance."'),
    ('Step 2 — Client Management / CRM (2 min)',
     'Navigate to Clients sidebar. Show client list with filters. Click Acme Corp: '
     'profile, billing info, revenue, associated projects, files, and members. Say: '
     '"We manage all client relationships here — from leads to active clients, with full financial tracking."'),
    ('Step 3 — Project Details (5 min)',
     'Click into "Acme Corp Website Redesign". Show: Overview (status, health, progress%), '
     'Tasks (Kanban — create and assign a task), Milestones, Timeline (Gantt), Updates (post one), '
     'Files (show visibility), Reports (show PDF), Financials (contract + invoice + Pay with Card).'),
    ('Step 4 — Developer View (2 min)',
     'Switch to Developer Dashboard. Show: My Open Tasks, Due Soon, In Review. '
     'Click a task → update status from IN_PROGRESS to REVIEW. Log 2 hours (billable). '
     'Say: "Developers have a clean, focused view — just their work, no clutter."'),
    ('Step 5 — QA / Findings (2 min)',
     'Navigate to Findings. Show filtered list by severity. Click a finding: '
     '"Login button not aligned on mobile" → Severity: High. Change status to In Review. '
     'View evidence screenshot. Show timeline. Say: "Our QA module tracks every issue with full audit history."'),
    ('Step 6 — Client Portal (2 min)',
     'Login as CLIENT_OWNER. Show Client Dashboard: only Acme Corp data visible. '
     'Show projects, client-visible updates, shared files, invoice with Pay with Card button. '
     'Say: "Clients get a professional portal — no access to internal data, no confusion."'),
    ('Step 7 — Settings & Integrations (1 min)',
     'Show Settings: Profile, 2FA setup, Notifications, Theme toggle. '
     'Show Admin Panel: Users, Roles, Invite links, Audit logs. '
     'Show Integrations: Slack, GitHub, Webhooks configured.'),
    ('Step 8 — Quick Feature Highlights (1 min)',
     'Ctrl+K → Global search (find anything instantly). '
     'Bell icon → Real-time notifications drawer. '
     'Language toggle → Arabic RTL layout. '
     'Automations → Show a rule: "When task status = DONE, notify PM".'),
]

for step_title, step_desc in demo_steps:
    add_heading(step_title, level=3)
    add_para(step_desc, size=10.5)
    doc.add_paragraph()

add_heading('Closing Statement', level=3)
add_para(
    '"Arena360 is not just a task manager — it\'s a complete delivery platform. '
    'Your PM runs projects, your developers deliver work, your finance team tracks revenue, '
    'and your clients stay informed — all in one system, with no friction."',
    italic=True, bold=True, size=11, color=ACCENT
)

# ═══════════════════════════════════════════════════════════════════
#  §8  KEY DIFFERENTIATORS
# ═══════════════════════════════════════════════════════════════════
add_heading('8. Key Differentiators vs Competitors', level=1)
add_table(
    ['#', 'Differentiator', 'Why It Matters'],
    [
        ['1', 'Client Portal + Financials in one',      'No need for a separate client communication or invoicing tool'],
        ['2', 'Findings / QA Module',                   'Dedicated issue tracker with severity, evidence, AI — rare in PM tools'],
        ['3', '4 Role-Based Dashboards',                'Each persona sees exactly what they need — PM, Dev, Finance, Client'],
        ['4', 'Multi-Tenancy + SSO (Google & SAML)',    'Multiple orgs, enterprise SSO — Google Workspace, Azure AD, Okta'],
        ['5', 'Stripe Payment Integration',             'Clients pay invoices directly in the platform — no external links'],
        ['6', 'Arabic / Full RTL Support',              'Purpose-built for MENA-region clients — rare competitive advantage'],
        ['7', 'Workflow Automation',                    'Multi-step approvals + rule-based triggers in one product'],
        ['8', 'Audit Trail (Before/After JSON)',        'Enterprise-grade compliance logging for every data change'],
    ],
    col_widths=[0.3, 2.3, 4.1]
)

# ═══════════════════════════════════════════════════════════════════
#  §9  COMPETITIVE COMPARISON
# ═══════════════════════════════════════════════════════════════════
add_heading('9. Competitive Comparison', level=1)
add_table(
    ['Feature', 'Arena360', 'Jira', 'Monday', 'Asana', 'ClickUp', 'Basecamp'],
    [
        ['Task Management',           '✅','✅','✅','✅','✅','✅'],
        ['Kanban Board',              '✅','✅','✅','✅','✅','❌'],
        ['Client Portal',             '✅','❌','✅','❌','❌','✅'],
        ['Financial Management',      '✅','❌','❌','❌','❌','❌'],
        ['Findings / QA Tracking',    '✅','✅','❌','❌','❌','❌'],
        ['Role-Based Dashboards',     '✅','❌','❌','❌','❌','❌'],
        ['Time Tracking',             '✅','✅','✅','✅','✅','✅'],
        ['Gantt / Timeline',          '✅','✅','✅','✅','✅','❌'],
        ['Automation',                '✅','✅','✅','✅','✅','❌'],
        ['Slack / GitHub Integration','✅','✅','✅','✅','✅','✅'],
        ['Google SSO / SAML',         '✅','✅','✅','✅','✅','✅'],
        ['Arabic / RTL',              '✅','✅','✅','❌','❌','❌'],
        ['Audit Trail',               '✅','✅','❌','❌','❌','❌'],
        ['PPTX / PDF Reports',        '✅','✅','✅','✅','✅','❌'],
        ['AI Features',               '✅','✅','✅','✅','✅','❌'],
        ['Stripe Payments (native)',   '✅','❌','❌','❌','❌','❌'],
    ],
    col_widths=[2.2, 0.85, 0.65, 0.85, 0.7, 0.75, 0.85]
)

add_para(
    'Overall Maturity Score: 7.6 / 10 — Enterprise-ready Mid-Market SaaS',
    bold=True, size=11, color=ACCENT
)

# ═══════════════════════════════════════════════════════════════════
#  §10  CLIENT PORTAL
# ═══════════════════════════════════════════════════════════════════
add_heading('10. Client Portal Experience', level=1)
add_para(
    'The client portal is a fully separate experience within the same application. '
    'Clients only see their own data — nothing internal, nothing from other clients.',
    size=11
)

add_heading('What Clients Can & Cannot See', level=2)
add_table(
    ['Section', 'CLIENT_OWNER', 'CLIENT_MANAGER', 'CLIENT_MEMBER'],
    [
        ['Their projects only',              '✅', '✅', '✅'],
        ['Client-visible status updates',    '✅', '✅', '✅'],
        ['Client-visible files',             '✅', '✅', '✅'],
        ['Their invoices',                   '✅', '❌', '❌'],
        ['Pay invoice via Stripe',           '✅', '❌', '❌'],
        ['Internal team discussions',        '❌', '❌', '❌'],
        ['Internal-only files',              '❌', '❌', '❌'],
        ['Other clients\' data',             '❌', '❌', '❌'],
    ],
    col_widths=[3.0, 1.2, 1.4, 1.2]
)

add_heading('Client Onboarding Steps', level=2)
onboarding = [
    'Admin creates Client account in the CRM module',
    'Admin adds a CLIENT_OWNER user → invite email sent automatically',
    'Client clicks invite link → sets their password',
    'Client logs in → sees their personalized client dashboard',
    'All communication through the platform: updates, files, discussions',
    'Invoice issued → Client notified → Client pays via Stripe (in-app)',
]
for s in onboarding: add_bullet(s)

# ═══════════════════════════════════════════════════════════════════
#  §11  DEPLOYMENT & INFRASTRUCTURE
# ═══════════════════════════════════════════════════════════════════
add_heading('11. Deployment & Infrastructure', level=1)
add_table(
    ['Aspect', 'Details'],
    [
        ['Local Dev',      'Frontend: npm run dev (port 5173) / Backend: npm run start:dev (port 3000)'],
        ['Docker',         'Full stack via npm run stack:up (API + PostgreSQL + MinIO)'],
        ['Database',       'PostgreSQL 15+; auto-migrations via Prisma'],
        ['Seeding',        'npx prisma db seed for initial demo/test data'],
        ['File Storage',   'MinIO S3-compatible (or local uploads/ folder for dev)'],
        ['API Docs',       'Swagger UI at http://localhost:3000/api-docs'],
        ['Health Check',   'GET /health endpoint always available'],
        ['Production',     'Docker Compose with .env.prod, restart policies, hidden ports'],
        ['Backup',         'scripts/backup-db.ps1 runs pg_dump; docs/backup-restore.md for DR procedure'],
        ['Rate Limiting',  '100 req/60s global limit; stricter on auth routes (5 req/min)'],
        ['Audit Logging',  'Before/after JSON snapshots of every mutation; sensitive data auto-redacted'],
        ['2FA',            'TOTP-based (Google Authenticator compatible); per-user setup/verify/disable'],
    ],
    col_widths=[1.8, 4.9]
)

# ═══════════════════════════════════════════════════════════════════
#  §12  FAQ
# ═══════════════════════════════════════════════════════════════════
add_heading('12. Frequently Asked Questions (FAQ)', level=1)

faqs = [
    (
        'Q: What makes Arena360 different from Jira or Monday.com?',
        'Arena360 combines project management with a built-in client portal, financial management '
        '(contracts, invoices, Stripe payments), and a dedicated QA/findings tracker — all in one '
        'product. Competitors require separate tools for each of these.'
    ),
    (
        'Q: Can clients log in and monitor their project progress?',
        'Yes. Clients get a dedicated portal login showing only their data — updates, files, and '
        'invoices. They can even pay invoices via Stripe directly in the app — no email chains needed.'
    ),
    (
        'Q: Does it support Arabic and right-to-left layout?',
        'Yes. Full Arabic (RTL) layout is supported with a one-click language toggle. All UI elements, '
        'navigation, forms, and data tables adapt to right-to-left reading direction.'
    ),
    (
        'Q: Is there Single Sign-On (SSO) support?',
        'Yes. Google OAuth 2.0 and SAML 2.0 are both supported for enterprise SSO, compatible '
        'with Okta, Microsoft Azure AD, Google Workspace, and any standard SAML Identity Provider.'
    ),
    (
        'Q: How secure is the platform?',
        'Arena360 includes: JWT authentication, bcrypt password hashing, TOTP-based 2FA, rate limiting '
        'on all endpoints, comprehensive audit logging (before/after JSON), org-level data isolation, '
        'and HTTPS enforcement.'
    ),
    (
        'Q: Can we generate professional reports for clients?',
        'Yes. The Reports tab generates PowerPoint (PPTX) and PDF reports from live project data. '
        'Reports go through a configurable approval workflow before being shared with clients.'
    ),
    (
        'Q: How does it handle multiple organizations / tenants?',
        'Arena360 is fully multi-tenant. Each organization has completely isolated data — multiple '
        'companies can use the same deployment with no data cross-contamination.'
    ),
    (
        'Q: What integrations are supported?',
        'Slack (team notifications), GitHub (create issues from findings), Stripe (invoice payments), '
        'Google SSO, SAML SSO, and custom webhooks for any external system integration.'
    ),
    (
        'Q: Are AI features available?',
        'Optional AI features via OpenAI API: project summaries, task suggestions, finding analysis '
        '(remediation and severity), status report generation, and contextual chat.'
    ),
    (
        'Q: How is the system backed up?',
        'A PowerShell backup script (scripts/backup-db.ps1) runs pg_dump. Full backup and restore '
        'documentation is included in docs/backup-restore.md for disaster recovery procedures.'
    ),
]

for q, a in faqs:
    add_para(q, bold=True, size=10.5, color=DARK_BLUE)
    add_para(a, size=10.5)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
#  FOOTER
# ═══════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('© Arena360  —  Confidential Client Presentation Document')
r.font.size = Pt(9)
r.font.color.rgb = MID_GRAY
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Version 1.0  |  March 2026  |  All rights reserved')
r.font.size = Pt(9)
r.font.color.rgb = MID_GRAY

# ─── Save ─────────────────────────────────────────────────────────
output_path = r'd:\Projects\a360\docs\Arena360_Presentation_Guide.docx'
doc.save(output_path)
print(f'✅ Document saved: {output_path}')
