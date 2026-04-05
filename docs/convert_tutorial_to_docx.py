
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configuration
MD_FILE = 'Client_Onboarding_Tutorial_EN_AR.md'
DOCX_FILE = 'Arena360_Client_Onboarding_Tutorial.docx'
IMAGE_DIR = 'images/tutorial/'

def create_tutorial_docx():
    if not os.path.exists(MD_FILE):
        print(f"Error: {MD_FILE} not found.")
        return

    doc = Document()
    
    # Title Page
    title = doc.add_heading('Arena360', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Client Onboarding & Handover Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo
    
    doc.add_paragraph('\n' * 2)
    
    # Read Markdown
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Handle Headers
        if line.startswith('######'):
            doc.add_heading(line[6:].strip(), level=6)
        elif line.startswith('#####'):
            doc.add_heading(line[5:].strip(), level=5)
        elif line.startswith('####'):
            doc.add_heading(line[4:].strip(), level=4)
        elif line.startswith('###'):
            doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith('##'):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith('#'):
            doc.add_heading(line[1:].strip(), level=1)

        # Handle Images ![alt](path)
        elif line.startswith('![') and '](' in line:
            match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
            if match:
                alt_text, img_path = match.groups()
                # Resolve path relative to doc folder
                # The MD says images/tutorial/dashboard.png
                # But we are in the doc folder already
                full_img_path = os.path.join(os.getcwd(), img_path)
                if os.path.exists(full_img_path):
                    doc.add_picture(full_img_path, width=Inches(6))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption if it's the next line or embedded in alt
                    caption = doc.add_paragraph(alt_text)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.runs[0].font.size = Pt(9)
                    caption.runs[0].italic = True
                else:
                    print(f"Warning: Image not found at {full_img_path}")

        # Handle Tables (Basic check)
        elif '|' in line and '-' not in line:
            # We skip complex table parsing for now or just treat as text
            # A real parser would be better, but we'll just add it as a paragraph
            p = doc.add_paragraph(line)
            if any("\u0600" <= c <= "\u06FF" for c in line): # Simple Arabic check
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Handle Bullet Points
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')

        # Regular Text
        else:
            p = doc.add_paragraph(line)
            # RTL Alignment for Arabic
            if any("\u0600" <= c <= "\u06FF" for c in line):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    doc.save(DOCX_FILE)
    print(f"Successfully generated {DOCX_FILE}")

if __name__ == '__main__':
    create_tutorial_docx()
